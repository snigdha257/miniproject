import base64
import os
import re
import tempfile
from datetime import datetime
from io import BytesIO
from typing import Optional

from bson import ObjectId
from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile, status
from fastapi.responses import Response
from pydantic import constr
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
import fitz

from database.connection import db
from database.schema import (
    count_documents_for_user,
    create_document,
    create_history_entry,
    create_mapping,
    get_document_for_user,
    get_documents_for_user,
    get_mapping_by_document_id,
    list_documents_for_user,
    update_document,
    update_mapping,
)
from routes.auth import get_current_user
from routes.schemas import (
    DashboardResponse,
    DocumentHistoryItem,
    DocumentUploadResponse,
    HistoryResponse,
    MaskRequest,
    MaskResponse,
    UnmaskRequest,
    UnmaskResponse,
    AnalyzeRequest,
    AnalyzeResponse,
)
from services.extractor import extract_text
from services.masking_engine import mask_text
from services.pipeline import detect_text_entities
from services.privacy_analyzer import analyze_privacy
from services.secure_masking import generate_key, secure_mask_text, unmask_text as secure_unmask_text, WrongKeyError
from services.groq_verifier import verify_with_groq

router = APIRouter(prefix="/documents", tags=["Documents"])

MAX_FILE_SIZE = 5 * 1024 * 1024
MAX_TEXT_LENGTH = 20000
PLATFORM_NAME = "Context-Aware Sensitive Data Masking"


def _safe_filename(document_name: Optional[str], fmt: str) -> str:
    if document_name:
        stem = os.path.splitext(document_name)[0] or "resume"
        return f"{stem}_masked.{fmt}"
    return f"resume_masked.{fmt}"


def _build_txt_bytes(text: str) -> bytes:
    return text.replace("\r\n", "\n").replace("\r", "\n").encode("utf-8")


def _build_docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    style = document.styles['Normal']
    style.font.name = 'Calibri' # type: ignore
    style.font.size = 11 # type: ignore

    paragraphs = re.split(r"\n(?:\s*\n)+", text.replace("\r\n", "\n").replace("\r", "\n"))
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            document.add_paragraph(paragraph_text)
        else:
            document.add_paragraph()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_plain_pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    story = []
    style = ParagraphStyle(name='Body', fontName='Helvetica', fontSize=10, leading=14, spaceAfter=6)
    paragraphs = re.split(r"\n(?:\s*\n)+", text.replace("\r\n", "\n").replace("\r", "\n"))
    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            story.append(Paragraph(paragraph_text, style))
        else:
            story.append(Spacer(1, 0.2 * inch))

    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.75 * inch, leftMargin=0.75 * inch, topMargin=0.75 * inch, bottomMargin=0.9 * inch)

    def add_footer(canvas_obj, doc_obj):
        canvas_obj.setFont('Helvetica', 8)
        canvas_obj.setFillColorRGB(0.45, 0.45, 0.45)
        canvas_obj.drawRightString(letter[0] - 0.75 * inch, 0.5 * inch, PLATFORM_NAME)
        canvas_obj.drawString(0.75 * inch, 0.5 * inch, f"Page {doc_obj.page}")

    doc.build(story, onLaterPages=add_footer, onFirstPage=add_footer)
    return buffer.getvalue()


def _build_redacted_pdf_bytes(file_bytes: bytes, entities: list[dict], masked_text: str) -> bytes:
    document = fitz.open(stream=file_bytes, filetype="pdf")
    for page in document:
        for entity in entities:
            entity_text = entity.get("text")
            if not isinstance(entity_text, str) or not entity_text.strip():
                continue
            for rect in page.search_for(entity_text):
                page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()

    buffer = BytesIO()
    document.save(buffer)
    document.close()
    return buffer.getvalue()


def _build_download_bytes(document: dict, fmt: str) -> tuple[bytes, str]:
    original_text = document.get("originalText") or document.get("raw_text") or ""
    masked_text = document.get("maskedText") or document.get("masked_text") or original_text or ""
    if not isinstance(masked_text, str):
        masked_text = str(masked_text)

    # Check if document has been restored (unmasked) - when maskedText equals originalText
    is_restored = masked_text == original_text and original_text

    if fmt == "txt":
        return _build_txt_bytes(masked_text), "text/plain; charset=utf-8"
    if fmt == "docx":
        return _build_docx_bytes(masked_text), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if fmt == "pdf":
        source_format = (document.get("source_format") or "").lower()
        file_bytes = document.get("file_bytes")
        if file_bytes and source_format == "pdf":
            # If restored (unmasked), return original PDF without redaction
            if is_restored:
                return file_bytes, "application/pdf"
            # If masked, return redacted PDF with entities masked
            return _build_redacted_pdf_bytes(file_bytes, document.get("entities", []), masked_text), "application/pdf"
        return _build_plain_pdf_bytes(masked_text), "application/pdf"
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format.")


@router.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    current_user: dict = Depends(get_current_user),
    file: Optional[UploadFile] = File(None),
    text: Optional[constr(strip_whitespace=True, min_length=1, max_length=MAX_TEXT_LENGTH)] = Form(None), # type: ignore
):
    if not file and not text:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Provide either a file or raw text.")

    filename = None
    source_format = None
    file_bytes = None
    if file:
        filename = file.filename
        contents = await file.read()
        file_bytes = contents
        if len(contents) == 0:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Uploaded file is empty.")
        if len(contents) > MAX_FILE_SIZE:
            raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="File too large.")

        suffix = os.path.splitext(filename)[1] or ".txt" # type: ignore
        source_format = suffix.lstrip(".").lower() if suffix else "txt"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(contents)
            temp_path = tmp_file.name
        try:
            raw_text = extract_text(temp_path)
        finally:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
    else:
        raw_text = text

    if not raw_text or not raw_text.strip():
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Document text is empty.")

    document = {
        "user_id": current_user["id"],
        "filename": filename,
        "originalText": raw_text,
        "maskedText": None,
        "riskScore": None,
        "mode": None,
        "privacy_report": {},
        "source_format": source_format,
        "file_bytes": file_bytes,
        "entities": [],
        "created_at": datetime.utcnow(),
    }
    document_id = create_document(document)
    return {"document_id": str(document_id)}


@router.post("/analyze", response_model=AnalyzeResponse)
async def analyze_document(request: AnalyzeRequest, current_user: dict = Depends(get_current_user)):
    try:
        document_id = ObjectId(request.document_id)
    except Exception:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid document ID.")

    document = get_document_for_user(document_id, current_user["id"])
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found.")

    original_text = document.get("originalText") or document.get("raw_text") or ""
    if not isinstance(original_text, str):
        original_text = str(original_text)

    local_entities = detect_text_entities(original_text)

    if request.processing_mode == "enhanced":
        verified_entities = verify_with_groq(original_text, local_entities)
    else:
        verified_entities = [
            {
                "text": ent["text"],
                "label": ent["label"],
                "recommendMask": True,
                "reason": "Detected by Standard Pipeline.",
                "start": ent["start"],
                "end": ent["end"]
            }
            for ent in local_entities
        ]

    update_document(document["_id"], {"entities": verified_entities})
    return {"entities": verified_entities}


@router.post("/mask", response_model=MaskResponse)
async def mask_document(request: MaskRequest, current_user: dict = Depends(get_current_user)):
    try:
        document_id = ObjectId(request.document_id)
    except Exception:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid document ID.")

    document = get_document_for_user(document_id, current_user["id"])
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found.")

    if request.mode == "standard":
        if request.style not in {"placeholder", "partial", "full"}:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid masking style for standard mode.")

        original_text = document.get("originalText") or document.get("raw_text") or ""
        if not isinstance(original_text, str):
            original_text = str(original_text)
        
        if request.entities is not None:
            entities = request.entities
        else:
            entities = detect_text_entities(original_text)

        privacy_report = analyze_privacy(entities)
        masked_result = mask_text(original_text, entities, request.style)
        update_data = {
            "mode": request.mode,
            "maskedText": masked_result["masked_text"],
            "riskScore": privacy_report["privacy_score"],
            "privacy_report": privacy_report,
            "entities": entities,
        }
        update_document(document["_id"], update_data)
        return {
            "masked_text": update_data["maskedText"],
            "privacy_score": privacy_report["privacy_score"],
            "risk_level": privacy_report["risk_level"],
            "entity_counts": privacy_report["entity_counts"],
        }

    if request.mode == "secure":
        if request.style != "placeholder":
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Secure mode only supports placeholder style.")
        key = generate_key()
        original_text = document.get("originalText") or document.get("raw_text") or ""
        if not isinstance(original_text, str):
            original_text = str(original_text)

        if request.entities is not None:
            entities = request.entities
        else:
            entities = detect_text_entities(original_text)

        masked_text, encrypted_blob = secure_mask_text(original_text, key, entities=entities)
        privacy_report = analyze_privacy(entities)
        encoded_blob = base64.b64encode(encrypted_blob).decode("utf-8")
        update_data = {
            "mode": request.mode,
            "maskedText": masked_text,
            "riskScore": privacy_report["privacy_score"],
            "privacy_report": privacy_report,
            "entities": entities,
        }
        update_document(document["_id"], update_data)
        update_mapping(str(document["_id"]), {"encryptedMapping": encoded_blob})
        create_history_entry({
            "user_id": current_user["id"],
            "document_id": str(document["_id"]),
            "date": datetime.utcnow(),
            "mode": request.mode,
        })
        return {
            "masked_text": masked_text,
            "privacy_score": privacy_report["privacy_score"],
            "risk_level": privacy_report["risk_level"],
            "entity_counts": privacy_report["entity_counts"],
            "secure_key": key.decode("utf-8"),
        }

    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid mode.")


@router.post("/unmask", response_model=UnmaskResponse)
async def unmask_document(request: UnmaskRequest, current_user: dict = Depends(get_current_user)):
    print(f"[DEBUG] Unmask request: document_id={request.document_id}, key={request.key[:20]}...")
    print(f"[DEBUG] Current user: {current_user.get('id')}")
    
    document = get_document_for_user(ObjectId(request.document_id), current_user["id"])
    print(f"[DEBUG] Retrieved document: {document is not None}")
    if document:
        print(f"[DEBUG] Document mode: {document.get('mode')}")
        print(f"[DEBUG] Document has maskedText: {bool(document.get('maskedText'))}")
    
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found.")
    
    mapping = get_mapping_by_document_id(str(document["_id"]))
    print(f"[DEBUG] Retrieved mapping: {mapping is not None}")
    if mapping:
        print(f"[DEBUG] Mapping has encryptedMapping: {bool(mapping.get('encryptedMapping'))}")
    
    if document.get("mode") != "secure" or not mapping or not mapping.get("encryptedMapping"):
        print(f"[DEBUG] Unmask failed - mode={document.get('mode')}, mapping={bool(mapping)}, encryptedMapping={bool(mapping.get('encryptedMapping')) if mapping else False}")
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Document is not stored in secure mode.")

    try:
        masked_text = document.get("maskedText") or document.get("masked_text") or ""
        if not isinstance(masked_text, str):
            masked_text = str(masked_text)
        print(f"[DEBUG] Masked text length: {len(masked_text)}")
        print(f"[DEBUG] Encrypted blob length: {len(mapping['encryptedMapping'])}")
        restored = secure_unmask_text(masked_text, base64.b64decode(mapping["encryptedMapping"]), request.key)
        print(f"[DEBUG] Restored text length: {len(restored)}")
    except WrongKeyError as exc:
        print(f"[DEBUG] WrongKeyError: {exc}")
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc))
    except Exception as exc:
        print(f"[DEBUG] Unexpected error: {type(exc).__name__}: {exc}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=str(exc))

    # Store the restored text back to the document so download returns the unmasked version
    update_document(document["_id"], {"maskedText": restored})

    return {"restored_text": restored}


@router.get("/dashboard", response_model=DashboardResponse)
async def dashboard(current_user: dict = Depends(get_current_user)):
    docs = get_documents_for_user(current_user["id"])
    documents_processed = len(docs)
    if documents_processed == 0:
        return {"documents_processed": 0, "average_privacy_score": 0.0, "entity_type_breakdown": {}}

    total_score = 0
    entity_type_breakdown = {}
    for doc in docs:
        privacy = doc.get("privacy_report", {})
        total_score += privacy.get("privacy_score", 0)
        for label, count in privacy.get("entity_counts", {}).items():
            entity_type_breakdown[label] = entity_type_breakdown.get(label, 0) + count

    return {
        "documents_processed": documents_processed,
        "average_privacy_score": round(total_score / documents_processed, 2),
        "entity_type_breakdown": entity_type_breakdown,
    }


@router.get("/{document_id}/download")
async def download_document(
    document_id: str,
    format: str = Query("txt", alias="format"),
    current_user: dict = Depends(get_current_user),
):
    try:
        object_id = ObjectId(document_id)
    except Exception:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid document ID.")

    document = get_document_for_user(object_id, current_user["id"])
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found.")

    fmt = (format or "txt").lower()
    if fmt not in {"txt", "docx", "pdf"}:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format.")

    payload, media_type = _build_download_bytes(document, fmt)
    filename = _safe_filename(document.get("filename"), fmt)
    headers = {"Content-Disposition": f"attachment; filename={filename}"}
    return Response(content=payload, media_type=media_type, headers=headers)


@router.get("/history", response_model=HistoryResponse)
async def history(
    current_user: dict = Depends(get_current_user),
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=50),
):
    skip = (page - 1) * page_size
    total = count_documents_for_user(current_user["id"])
    docs = list_documents_for_user(current_user["id"], skip=skip, limit=page_size)

    documents = []
    for doc in docs:
        documents.append(
            {
                "document_id": str(doc["_id"]),
                "filename": doc.get("filename"),
                "created_at": doc.get("created_at"),
                "mode": doc.get("mode") or "standard",
                "risk_level": doc.get("privacy_report", {}).get("risk_level", "LOW"),
            }
        )

    return {"documents": documents, "page": page, "page_size": page_size, "total": total}
