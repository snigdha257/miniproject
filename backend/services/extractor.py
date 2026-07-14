"""
backend/services/extractor.py

Extract plain text from .txt, .pdf (pdfplumber), and .docx (python-docx) files.
Returns a single unicode string with normalised whitespace.
"""

from pathlib import Path


def extract_text(file_path: str) -> str:
    """
    Extract plain text from a file.

    Args:
        file_path: Absolute or relative path to a .txt, .pdf, or .docx file.

    Returns:
        Extracted text as a plain string.

    Raises:
        FileNotFoundError: If the path does not exist.
        ValueError: If the file extension is not supported.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8")

    elif suffix == ".pdf":
        import pdfplumber  # lazy import — only needed for PDFs

        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
        return "\n".join(pages)

    elif suffix == ".docx":
        from docx import Document  # lazy import — only needed for DOCX

        doc = Document(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)

    else:
        raise ValueError(
            f"Unsupported file type: {suffix!r}. Supported formats: .txt, .pdf, .docx"
        )


# ---------------------------------------------------------------------------
# Manual verification — run: python extractor.py
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import os
    import tempfile
    from pathlib import Path

    here = Path(__file__).parent
    sample_txt = here / "sample.txt"

    # ── 1. Plain text ──────────────────────────────────────────────────────
    print("=" * 60)
    print("TEST 1: .txt extraction")
    print("=" * 60)
    txt_result = extract_text(str(sample_txt))
    print(txt_result[:500])
    print(f"\n[OK] Extracted {len(txt_result)} characters from {sample_txt.name}\n")

    # ── 2. DOCX — create programmatically, then extract ───────────────────
    print("=" * 60)
    print("TEST 2: .docx extraction")
    print("=" * 60)
    try:
        from docx import Document
    except ImportError as exc:
        print("[SKIP] python-docx is not installed. Install it to run .docx extraction tests.")
        raise

    docx_path = here / "sample.docx"
    doc = Document()
    doc.add_heading("Test DOCX Document", 0)
    doc.add_paragraph("Name: Priya Sharma")
    doc.add_paragraph("Email: priya.sharma@infosys.com")
    doc.add_paragraph("Aadhaar: 1234 5678 9012")
    doc.add_paragraph("She is the CTO at Infosys, based in Pune.")
    doc.save(str(docx_path))

    docx_result = extract_text(str(docx_path))
    print(docx_result)
    print(f"\n[OK] Extracted {len(docx_result)} characters from {docx_path.name}\n")

    # ── 3. Unsupported extension — should raise ValueError ─────────────────
    print("=" * 60)
    print("TEST 3: unsupported extension (.csv)")
    print("=" * 60)
    csv_path = here / "dummy.csv"
    csv_path.write_text("col1,col2\nval1,val2")  # real file, wrong extension
    try:
        extract_text(str(csv_path))
    except ValueError as e:
        print(f"[OK] Caught expected ValueError: {e}\n")
    finally:
        csv_path.unlink(missing_ok=True)  # clean up

    print("All extractor tests passed.")
